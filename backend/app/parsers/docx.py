"""Parse an email-like item exported to .docx into an EmailParseResult.

Only Office Open XML (.docx) is supported; legacy OLE2 .doc is out of scope and
never reaches here (the suffix dispatch in ingest only routes .docx). Documents
whose text does not open with an email header cluster raise NotAnEmailError
(handled by ingest as a skipped, non-fatal file).
"""

from io import BytesIO
from typing import Iterator

from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from .base import EmailParseResult, EmptyDocumentError
from ..services.email_text import build_printed_email_result


def _iter_block_items(parent) -> Iterator[object]:
    """Yield Paragraph and Table objects from a document body in document order.

    python-docx exposes document.paragraphs and document.tables separately, which
    loses ordering. Walking the underlying XML body children preserves the order
    so a header table at the top stays at the top. Standard python-docx recipe.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("unsupported parent type")
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _table_text(table: Table) -> str:
    lines = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        lines.append("\t".join(c for c in cells if c))
    return "\n".join(line for line in lines if line)


def _extract_text(content: bytes) -> str:
    doc = Document(BytesIO(content))
    parts = []
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            parts.append(block.text)
        elif isinstance(block, Table):
            parts.append(_table_text(block))
    return "\n".join(parts).strip()


def parse_docx_bytes(
    content: bytes, source_file_display: str, default_tz: str = "UTC"
) -> EmailParseResult:
    text = _extract_text(content)
    if not text:
        raise EmptyDocumentError("No extractable text in .docx document.")
    return build_printed_email_result(
        text, source_file_display, default_tz, source_format="docx"
    )
